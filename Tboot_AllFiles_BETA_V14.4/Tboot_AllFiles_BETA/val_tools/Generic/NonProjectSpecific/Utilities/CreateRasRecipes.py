#!/usr/bin/env python
#+----------------------------------------------------------------------------+
#| INTEL CONFIDENTIAL
#| Copyright 2014 Intel Corporation All Rights Reserved.
#| 
#| The source code contained or described herein and all documents related
#| to the source code ("Material") are owned by Intel Corporation or its
#| suppliers or licensors. Title to the Material remains with Intel Corp-
#| oration or its suppliers and licensors. The Material may contain trade
#| secrets and proprietary and confidential information of Intel Corpor-
#| ation and its suppliers and licensors, and is protected by worldwide
#| copyright and trade secret laws and treaty provisions. No part of the
#| Material may be used, copied, reproduced, modified, published, uploaded,
#| posted, transmitted, distributed, or disclosed in any way without
#| Intel's prior express written permission.
#| 
#| No license under any patent, copyright, trade secret or other intellect-
#| ual property right is granted to or conferred upon you by disclosure or
#| delivery of the Materials, either expressly, by implication, inducement,
#| estoppel or otherwise. Any license under such intellectual property
#| rights must be express and approved by Intel in writing.
#+----------------------------------------------------------------------------+
#| $Id: CreateRasRecipes.py 182 2015-06-18 21:31:49Z amr\egross $
#| $Date: 2015-06-18 14:31:49 -0700 (Thu, 18 Jun 2015) $
#| $Author: amr\egross $
#| $Revision: 182 $
#+----------------------------------------------------------------------------+
#| TODO:
#|   *  Fill in the script description
#|   *  Clean up processSpreadsheetRow
#|   *  Command line option for Template File
#|   *  Command line option for RAS Features Spreadsheet
#|   *  Command line option to create recipe from only single row
#|   *  Command line option to stop after the first recipe
#|   *  Add check to make sure we actually found all specified template tokens
#+----------------------------------------------------------------------------+

"""
    Write something here that summarizes what this script does
"""
# Standard Libary Imports
import os           as _os
import sys          as _sys
import re           as _re
import logging      as _logging
from optparse import OptionParser

# Downloaded Library Imports
from docx import Document
from datetime import date
from openpyxl import load_workbook

## Global Variables/Constants
bDebug                  = False
nOutputWidth            = 80
__version__             = "$Rev: 182 $".replace("$Rev:","").replace("$","").strip()
sScriptName             = _re.split("\.", _os.path.basename(__file__))[0]
sLogfileName            = '%s_pid%d.log' % (sScriptName, _os.getpid())

# val_tools Utilities Import - gotta find it first!
sScriptPath = _os.path.dirname(__file__)
if (bDebug): 
    print "ScriptPath:                  %s" % sScriptPath
sUtilitiesPath = sScriptPath + "."  #  <--- make sure this is the correct relative path!
if (bDebug): 
    print "ValToolsUtilsPath:           %s" % sUtilitiesPath
sUtilitiesPath =  _os.path.normpath(sUtilitiesPath)
if (bDebug):
    print "NormalizedValToolsUtilsPath: %s" % sUtilitiesPath
_sys.path.append(sUtilitiesPath)
import ValToolsUtilities as _ValToolsUtilities

#  Since we may want to import functionality from this script into another script,
#  only create the Logger instance if this is executing as a script and not being
#  imported as a module
if __name__ == '__main__':
    lLogger = _ValToolsUtilities.setupLogger(bDebug, sLogfileName)

# Dictionary of Feature Data that we want to extract from the XLSX file
# Used to store column number and other relevant data for processing 
# feature information
dictFeatures = {
    'Category'                                : {'ColNum' : None, 'Found': False, 'SpecialProcessing': False,  },
    'Feature SN'                              : {'ColNum' : None, 'Found': False, 'SpecialProcessing': False,  },
    'Purley Requirement HSD'                  : {'ColNum' : None, 'Found': False, 'SpecialProcessing': True,   },
    'POR / NOT POR / SOW'                     : {'ColNum' : None, 'Found': False, 'SpecialProcessing': False,  },
    'Platform RAS Feature'                    : {'ColNum' : None, 'Found': False, 'SpecialProcessing': False,  },
    'External Description'                    : {'ColNum' : None, 'Found': False, 'SpecialProcessing': False,  },
    'Revision'                                : {'ColNum' : None, 'Found': False, 'SpecialProcessing': True,   },
    'Recipe Intro'                            : {'ColNum' : None, 'Found': False, 'SpecialProcessing': False,  },
    'Use Cases'                               : {'ColNum' : None, 'Found': False, 'SpecialProcessing': False,  },
    'Feature Coverage'                        : {'ColNum' : None, 'Found': False, 'SpecialProcessing': False,  },
    'Recipe1 Steps'                           : {'ColNum' : None, 'Found': False, 'SpecialProcessing': False,  },
    'Recipe2 Steps'                           : {'ColNum' : None, 'Found': False, 'SpecialProcessing': False,  },
    'BKC'                                     : {'ColNum' : None, 'Found': False, 'SpecialProcessing': False,  },
    'Feature Gap List'                        : {'ColNum' : None, 'Found': False, 'SpecialProcessing': False,  },
    'Validation Tools: Setup'                 : {'ColNum' : None, 'Found': False, 'SpecialProcessing': False,  },
    'Validation Tools: System Stress'         : {'ColNum' : None, 'Found': False, 'SpecialProcessing': False,  },
    'Validation Tools: Induce Error'          : {'ColNum' : None, 'Found': False, 'SpecialProcessing': False,  },
    'Validation Tools: Verify Results'        : {'ColNum' : None, 'Found': False, 'SpecialProcessing': False,  },
    'Validation Requirements: Setup'          : {'ColNum' : None, 'Found': False, 'SpecialProcessing': False,  },
    'Validation Requirements: System Stress'  : {'ColNum' : None, 'Found': False, 'SpecialProcessing': False,  },
    'Validation Requirements: Induce Error'   : {'ColNum' : None, 'Found': False, 'SpecialProcessing': False,  },
    'Validation Requirements: Verify Results' : {'ColNum' : None, 'Found': False, 'SpecialProcessing': False,  },
    'Customer Visibility'                     : {'ColNum' : None, 'Found': False, 'SpecialProcessing': False,  },
    'Can Customer Enable/Disable Feature?'    : {'ColNum' : None, 'Found': False, 'SpecialProcessing': False,  },
    'Reference Docs'                          : {'ColNum' : None, 'Found': False, 'SpecialProcessing': False,  },
}


# Base URL for Purley HSD Requirements (append req number for full URL)
sPurleyBaseUrl = 'https://vthsd.fm.intel.com/hsd/purley_platform/requirement/default.aspx?requirement_id='

# Dictionary of characters that are illigal to use in filename
# and their replacements to use in creating filenames
dictReplaceIllegalFilenameChars = {
    "\n": " ",
    ":" : "-",
    ">" : "-",
    "<" : "-",
    "/" : "-",
    "?" : "-",
    '"' : "'",
}

#+----------------------------------------------------------------------------+
#|  Handle Command Line Options
#|
#|  This functon defines all supported command line options and invokes the
#|  methods used to extract those options from the user-supplied command line
#|
#|  Inputs:     None
#|  Returns:    Command Line Options Object from OptionParser
#|
#+----------------------------------------------------------------------------+
def parseCommandLine():

    #  Create a parser object and add options to it
    parser = OptionParser()
    parser.add_option("--debug", action="store_true",
                      dest="Debug", default=False,
                      help="Turn on DEBUG functionality of script.")

    parser.add_option("--recipenum", action="store", dest="RecipeNum", 
                      type="int", default=-1,
                      help="Specific recipe to be generated (row number in XLSX input file).  Default of -1 means ALL rows.")

    parser.add_option("--docx_template_location", action="store", dest="DocxTemplateLocation", 
                      default='//dpsvlx03.dp.intel.com/lab/Work/egross/RasRecipes/MasterSource/Template_ValidationRecipe.docx',
                      help="Location of template file for the DOCX output of the script. Default is //dpsvlx03.dp.intel.com/lab/Work/egross/RasRecipes/MasterSource/Template_ValidationRecipe.docx")

    parser.add_option("--xlsx_input_file_location", action="store", dest="XlsxInputLocation", 
                      default='//dpsvlx03.dp.intel.com/lab/Work/egross/RasRecipes/MasterSource/Purley RAS POR - RAS CoE Coverage Matrix.xlsx',
                      help="Location of input Feature/Recipe data in XLSX format. Default is //dpsvlx03.dp.intel.com/lab/Work/egross/RasRecipes/MasterSource/Purley RAS POR - RAS CoE Coverage Matrix.xlsx")


    #  Process the actual command line and split it into options and arguments
    (oCmdlineOptions, aArgs) = parser.parse_args()

    #  Set global bDebug variable and logger mesaging level if necessary
    if (oCmdlineOptions.Debug):
        global bDebug
        bDebug = oCmdlineOptions.Debug
        lLogger.setLevel(_logging.DEBUG)

    #  Debug output to indicate what the results of command line processing are
    lLogger.debug("Debug                Option read as %s"  % oCmdlineOptions.Debug        )
    lLogger.debug("RecipeNum            Option read as %s" % oCmdlineOptions.RecipeNum     )
    lLogger.debug("DocxTemplateLocation Option read as %s" % oCmdlineOptions.RecipeNum     )
    lLogger.debug("XlsxInputLocation    Option read as %s" % oCmdlineOptions.RecipeNum     )

    #  Return options data structure
    return oCmdlineOptions

#+----------------------------------------------------------------------------+
#| Print a divider to the screen
#|
#|  Inputs:     
#|              [optional] Character to use
#|              [optional] Indentation (number of spaces) for output
#|              [optional] Length of output
#|              [optional] Type of output (e.g. info, debug, error)
#|
#|  Returns:    True on success; False otherwise
#|
#+----------------------------------------------------------------------------+
def printDivider (sCharacter='-', nIndent=0, nLength=nOutputWidth, sInfoLevel='debug'):
    if   (sInfoLevel == 'debug') :
        lLogger.debug("%s%s" % (' ' * nIndent, sCharacter * nLength))
    elif (sInfoLevel == 'info') :
        lLogger.info ("%s%s" % (' ' * nIndent, sCharacter * nLength))
    elif (sInfoLevel == 'error') :
        lLogger.error("%s%s" % (' ' * nIndent, sCharacter * nLength))
    else:
        return False
    return True


#+----------------------------------------------------------------------------+
#| Function to take all characters in input text that are illegal to use in
#| filenames and replace them with substitute printable characters
#|
#|  Inputs:     Text (string)
#|  Returns:    Modified Text (string)
#|
#+----------------------------------------------------------------------------+
def replaceIllegalFilenameCharacters (sText):
    for sBadChar in dictReplaceIllegalFilenameChars.keys():
        sText = sText.replace(sBadChar, dictReplaceIllegalFilenameChars[sBadChar])
    return sText



#+----------------------------------------------------------------------------+
#| Given a spreadsheet row, search for all the data we're interested in and
#| note the column number that contains that information
#|
#|  Inputs:     
#|              Spreasheet Row object containing header information
#|              Dictionary storing locations of data of interest
#|
#|  Returns:    True upon success; False otherwise
#|
#+----------------------------------------------------------------------------+
def processHeaderRow (wsrHeaderInfo, dictFeatures):
    lLogger.info("Second row has header data.  Extracting...")
    nColumn = 0
    # Loop through all the cells in the row and look for headers matching
    # the data we're interested in.  If it matches, record the column number
    # for use later
    for wscFeatureData in wsrHeaderInfo:
        sCellData = "%s" % wscFeatureData.value
        if (sCellData in dictFeatures.keys()):
            dictFeatures[sCellData]['ColNum'] = nColumn
        nColumn = nColumn + 1

    # Summarize all the data we've found
    printDivider(nIndent=4, nLength=nOutputWidth-4)
    lLogger.debug("    Found Feature Data Types In These Header Column Numbers:")
    for dataType in sorted(dictFeatures.keys()):
        if (dictFeatures[dataType]['ColNum'] == None) :
            lLogger.error("ERROR:  didn't find a column with '%s' data in it." % dataType)
            lLogger.error("    I need this to continue.  Please fix the spreadsheet.")
            exit(1)
        else:
            lLogger.debug("    %30s : %3d" % (dataType, dictFeatures[dataType]['ColNum']))
    return True

#+----------------------------------------------------------------------------+
#| Given a spreadsheet row, extract all data of interest based on information
#| obtained from the header row
#|
#|  Inputs:     
#|              Spreasheet Row object containing feature data
#|              Dictionary storing locations of data of interest
#|              Dictionary storing data extracted for current feature
#|
#|  Returns:    True upon success; False otherwise
#|
#+----------------------------------------------------------------------------+
def processDataRow (wsrDataInfo, dictFeatures, dictCurFeatureData):

    # Extract the data from the current row for the columns we're
    # interested in
    nColumn = 0
    for sDataType in dictFeatures.keys():
        nDataColumn = dictFeatures[sDataType]['ColNum']
        dictCurFeatureData[sDataType] = dict()
        dictCurFeatureData[sDataType]['cellvalue'] = wsrDataInfo[nDataColumn].value

    # Print out what we found
    lLogger.debug("    Found Feature Data:")
    printDivider(nIndent=4, nLength=nOutputWidth-4)
    for sDataType in sorted(dictCurFeatureData.keys()):
        sText = "%s" % dictCurFeatureData[sDataType]['cellvalue']
        sText = sText.encode('utf-8')
        sText = sText.replace("\n", " ")
        lLogger.debug("    %30s : %s" % (sDataType, sText))
    printDivider(nIndent=4, nLength=nOutputWidth-4)

    return True

#+----------------------------------------------------------------------------+
#| Create the Title for the document and strip out any newlines
#|
#|  Inputs:     
#|              Dictionary storing data extracted for current feature
#|
#|  Returns:    String containing Title for document
#|
#+----------------------------------------------------------------------------+
def createDocTitle (dictCurFeatureData):
    sDocTitle = "%s RAS (%s) : %s" % (
                                        dictCurFeatureData['Category']['cellvalue'],
                                        dictCurFeatureData['Feature SN']['cellvalue'],
                                        dictCurFeatureData['Platform RAS Feature']['cellvalue'],
                                     )
    sDocTitle  = sDocTitle.replace("\n", " ")
    return sDocTitle


#+----------------------------------------------------------------------------+
#| Using data gathered from spreadsheet row, decide if this is a feature of
#| interest and create a recipe document if appropriate
#|
#|  Inputs:     
#|              Dictionary storing data extracted for current feature
#|
#|  Returns:    True upon success; False otherwise
#|
#+----------------------------------------------------------------------------+
def processSpreadsheetRow (dictCurFeatureData, sDocxTemplateLocation):
    bErrorsOccurred = False

    # Create the text title for the document based on the feature data
    sDocTitle = createDocTitle(dictCurFeatureData)

    # If the "Category" Column is empty, then this is not valid data; move on.
    if (dictCurFeatureData['Category']['cellvalue'] == None):
        lLogger.info("Blank Row")

    # If this is not a POR feature, move on
    elif not (dictCurFeatureData['POR / NOT POR / SOW']['cellvalue'] == 'POR'):
        lLogger.info("Non-POR Feature: %s" % sDocTitle.encode('utf-8') )

    # This is a valid row for feature data - process it!
    else:

        lLogger.info("    Title    : %s" % sDocTitle.encode('utf-8'))

        # Create the filename for the output document and get rid of
        # any characters that aren't allowed in Windows filenames
        sDocFileName  = "Purley %s%s" % (sDocTitle, ".docx")
        sDocFileName = replaceIllegalFilenameCharacters(sDocFileName)
        lLogger.info("    Filename : %s" % sDocFileName.encode('utf-8'))

        #docTemplate = Document('../Template_ValidationRecipe.docx')
        docTemplate = Document(sDocxTemplateLocation)

        #import urllib
        #testfile = urllib.URLopener()
        #testfile.retrieve("https://soco.intel.com/servlet/JiveServlet/downloadBody/1960346-102-9-654301/Template_ValidationRecipe.docx", "Template_ValidationRecipe.docx")

        # Modify the template by inserting new data
        for paraInputParagraph in docTemplate.paragraphs:
            # Paragraphs tagged with DELETEME should be removed
            if 'DELETEME:' in paraInputParagraph.text:
                lLogger.debug("    Removing template-only paragraph")
                paraInputParagraph.text = ""


            # Look for tags in the text indicating where we should insert things
            else: 
                bFoundTokenMatch = False
                for sFeatureDictKey in dictCurFeatureData.keys():
                    sDocxToken = "[INSERT %s HERE]" % sFeatureDictKey
                    if sDocxToken in paraInputParagraph.text:
                        if (dictFeatures[sFeatureDictKey]['SpecialProcessing'] == False):
                            dictFeatures[sFeatureDictKey]["Found"] = True
                            sParagraph = dictCurFeatureData[sFeatureDictKey]['cellvalue']
                            lLogger.debug("%s %35s    Match: %s" % (" " * 40, sFeatureDictKey, paraInputParagraph.text.encode('utf-8')))
                            lLogger.debug("    Paragraph: %s    " % (sParagraph))
                            paraInputParagraph.text = sParagraph
                        elif(sFeatureDictKey == 'Revision'):
                            dictFeatures[sFeatureDictKey]["Found"] = True
                            sParagraph = "v%s" % str(dictCurFeatureData[sFeatureDictKey]['cellvalue'])
                            lLogger.debug("%s %35s    Match: %s" % (" " * 40, sFeatureDictKey, paraInputParagraph.text.encode('utf-8')))
                            #lLogger.debug("    Paragraph: %s    " % (sParagraph))
                            paraInputParagraph.text = sParagraph
                        elif(sFeatureDictKey == 'Purley Requirement HSD'):
                            dictFeatures[sFeatureDictKey]["Found"] = True
                            sPurleyHsdNum = dictCurFeatureData['Purley Requirement HSD']['cellvalue']
                            lLogger.debug("%s %35s    Match: %s" % (" " * 40, sFeatureDictKey, paraInputParagraph.text.encode('utf-8')))
                            paraInputParagraph.text = ""
                            paraInputParagraph.add_run("Purley HSD Number: ") 
                            paraInputParagraph.add_run(str(sPurleyHsdNum)).bold=True
                            paraInputParagraph.add_run("\n") 
                            paraInputParagraph.add_run("\n") 
                            paraInputParagraph.add_run("Purley URL: %s%s)" % (sPurleyBaseUrl, sPurleyHsdNum))
             
                        else:
                            lLogger.error("Found token that requires special script processing, but")
                            lLogger.error("    there is no code to process that token.  Need to update this script!")

                if not bFoundTokenMatch:
                    sData = "Date"
                    sDocxToken = "[INSERT %s HERE]" % sData
                    if sDocxToken in paraInputParagraph.text:
                        dateToday = date.today()
                        sParagraph = dateToday.strftime("%B %d, %Y")
                        lLogger.debug("%s %35s    Match: %s" % (" " * 40, sData, paraInputParagraph.text.encode('utf-8')))
                        paraInputParagraph.text = sParagraph

                    sData = "Title"
                    sDocxToken = "[INSERT %s HERE]" % sData
                    if sDocxToken in paraInputParagraph.text:
                        sParagraph = sDocTitle
                        lLogger.debug("%s %35s    Match: %s" % (" " * 40, sData, paraInputParagraph.text.encode('utf-8')))
                        paraInputParagraph.text = sParagraph


#           elif '[InsertRevisionHere]' in paraInputParagraph.text:
#               sRevision = dictCurFeatureData['Revision']['cellvalue']
#               lLogger.debug("    Revision Match: %s (%s)" % (paraInputParagraph.text.encode('utf-8'), sRevision))
#               paraInputParagraph.text = "v%0.2f" % (float(sRevision))

        # When we're done with the row, then we need to save the file
        docTemplate.save(sDocFileName)


        for sKey in dictFeatures.keys():
            #  report any tokens not found after we're done with the template
            if (dictFeatures[sKey]["Found"] == False):
                lLogger.error("After processing template, '%s' token was not found.  Data not imported." % sKey)
                bErrorsOccurred = True
            #  reset indicator in preparation for next time through the template
            else:
                dictFeatures[sKey]["Found"] == False


    return not bErrorsOccurred


#+----------------------------------------------------------------------------+
#| Process each row of the input spreadsheet, creating Recipe documents as
#| necessary
#|
#|  Inputs:     
#|              None
#|
#|  Returns:    True upon success; False otherwise
#|
#+----------------------------------------------------------------------------+
def createRecipes (nRecipeNum, sDocxTemplateLocation, sXlsxInputLocation):
    # Load the workbook and go to the 'Coverage Mapping' sheet
    wbRasFeatures = load_workbook(filename=sXlsxInputLocation)
    #wbRasFeatures = load_workbook(filename="../MasterRASFeatures.xlsx")
    wsCoverageMapping = wbRasFeatures['Coverage Mapping']
    
    # Initialize some global variables
    nExcelRow = 1  # Excel starts row numbering with 1
    
    # Loop through all the rows in the worksheet, looking for feature data
    for wsrInput in wsCoverageMapping.rows:

        # First row expected to be blank
        if (nExcelRow == 1) :
            printDivider(sInfoLevel='info', sCharacter='=')
            lLogger.info("First row expected to be blank")
        # Second row is header row - use this to find the columns containing
        # the data we're interested in
        elif (nExcelRow == 2):
            printDivider(sInfoLevel='info', sCharacter='=')
            processHeaderRow(wsrInput, dictFeatures)
    
        # All valid rows beyond the first two should be valid feature data
        else:
            if (nRecipeNum == -1) or (nExcelRow == nRecipeNum):
                printDivider(sInfoLevel='info', sCharacter='=')
                lLogger.debug("Valid rows beyond the second row should have feature data... Extracting Row%03d" % nExcelRow)
                lLogger.debug("Row meets processing criteria; creating recipe document...")
                # Create a data structure for holding the data related to the feature
                # defined in the row we're currently processing
                dictCurFeatureData = dict()
                processDataRow(wsrInput, dictFeatures, dictCurFeatureData)
    
    
                # Using data gathered from spreadsheet row, decide if this is a feature of
                # interest and create a recipe document if appropriate
                processSpreadsheetRow(dictCurFeatureData, sDocxTemplateLocation)
            else:
                printDivider(sInfoLevel='debug', sCharacter='=')
                lLogger.debug("Row did not meet processing criteria; NOT creating recipe document...")
    
        # Increment our row counter because we're now gonna parse the next row
        nExcelRow = nExcelRow + 1

    return True

#+----------------------------------------------------------------------------+
#    XX     XX  XXX      XXXXX  XX   XXXX
#     X     X     X        X     X     X
#     XX   XX    X X       X     XX    X
#     XX   XX    X X       X     X X   X
#     X X X X   X   X      X     X X   X
#     X X X X   X   X      X     X  X  X
#     X  X  X   XXXXX      X     X   X X
#     X  X  X  X     X     X     X   X X
#     X     X  X     X     X     X    XX
#    XXX   XXXXXX   XXX  XXXXX  XXXX   X
#+----------------------------------------------------------------------------+
def main():
    #  Variable definitions
    bErrorsOccurred = False # used to short-circuit certain steps if errors found

    #  Startup tasks - get the logger configured
    _ValToolsUtilities.printStartupBanner(lLogger, nOutputWidth, 
                                          sScriptName, __version__)

    #  Get command line options, if any
    oCmdlineOptions = parseCommandLine()


    #  Figure out what recipes we need to create and create them!
    bErrorsOccurred = not createRecipes(
                                            oCmdlineOptions.RecipeNum,
                                            oCmdlineOptions.DocxTemplateLocation,
                                            oCmdlineOptions.XlsxInputLocation,
                                       )

    #  Return boolean indicating whether we were successful or not
    _ValToolsUtilities.printFinishingBanner(lLogger, bErrorsOccurred, nOutputWidth,
                                            sScriptName, __version__)
    return (not bErrorsOccurred)
    

####################################################################################

if __name__ == '__main__':
    if main():
        lLogger.info("Exiting with zero status...")
        _sys.exit(0)  # zero exit status means script completed successfully
    else:
        lLogger.error("Exiting with non-zero status...")
        _sys.exit(1)  # non-zero exit status means script did not complete successfully





# Code to attempt adding a table to the document (WIP)
#                            table = docTemplate.add_table(rows=1, cols=3)
#                            hdr_cells = table.rows[0].cells
#                            hdr_cells[0].text = 'Item'
#                            hdr_cells[1].text = 'Data'
#                            for sItem in ('Feature SN', 'Category', 'Purley Requirement HSD'):
#                                row_cells = table.add_row().cells
#                                row_cells[0].text = str(sItem)
#                                row_cells[1].text = str(dictCurFeatureData[sItem]['cellvalue'])
